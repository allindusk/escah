#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将 tools/_texts_for_translation 下的所有 .txt 转换为 .docx（多线程）。

- 镜像子目录结构输出到 tools/_docx_out/（characters/ 子目录同样保留）
- 每个 txt 一行 = docx 一个段落，完整保留 [N] 原文
- 设置中日文友好东亚字体，防止日文显示为方块
- 使用 ThreadPoolExecutor 并发处理（IO/序列化密集，线程足够）
"""
import os
import glob
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC_ROOT = os.path.join(os.path.dirname(__file__), "_texts_for_translation")
OUT_ROOT = os.path.join(os.path.dirname(__file__), "_docx_out")

# 东亚字体：保证日文（含汉字/假名）正常显示
EA_FONT = "Yu Gothic"


def set_east_asian_font(style):
    """给 Normal 样式设置东亚字体。"""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), EA_FONT)
    rfonts.set(qn("w:ascii"), EA_FONT)
    rfonts.set(qn("w:hAnsi"), EA_FONT)


def convert_one(txt_path: str) -> tuple:
    """转换单个 txt -> docx，返回 (src, out, n_lines)。"""
    rel = os.path.relpath(txt_path, SRC_ROOT)
    out_path = os.path.join(OUT_ROOT, rel[:-4] + ".docx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    with open(txt_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    set_east_asian_font(doc.styles["Normal"])
    # 文件名作为标题
    doc.add_heading(os.path.basename(txt_path)[:-4], level=1)
    for line in lines:
        doc.add_paragraph(line)

    doc.save(out_path)
    return txt_path, out_path, len(lines)


def main():
    files = []
    for p in glob.glob(os.path.join(SRC_ROOT, "**", "*.txt"), recursive=True):
        files.append(p)
    print(f"发现 {len(files)} 个 txt 文件，开始多线程转换 -> {OUT_ROOT}")

    max_workers = max(1, round((os.cpu_count() or 4) * 0.8))  # 留约 20% CPU 不占满
    total_lines = 0
    done = 0
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        futures = {ex.submit(convert_one, p): p for p in files}
        for fut in as_completed(futures):
            src, out, n = fut.result()
            total_lines += n
            done += 1
            if done % 50 == 0 or done == len(files):
                print(f"  [{done}/{len(files)}] {os.path.basename(src)} -> {n} 行")

    print(f"完成：{len(files)} 个文件，共 {total_lines} 行，输出在 {OUT_ROOT}")


if __name__ == "__main__":
    main()
