#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""合并 / 拆分 / 转 docx：characters 待翻译文本（独立工具，不影响原项目流程）。

合并 (merge)：把 tools/_texts_for_translation/characters/*.txt 的 369 个文件拼成一个文件。
      每个文件用纯大写字母+数字的标记 `A001` ... `A369` 分隔（不含任何可翻译单词）。
      序号 -> 原日文文件名的映射单独存到 manifest.txt（不随翻译文件发出，仅供拆分还原）。

转 docx (docx)：把 characters_merged.txt 转成 characters_merged.docx（每行一段落，
      A001 等标记行加粗提示勿译），供在 Word 里翻译。

拆分 (split)：读取合并文件（.txt 或 .docx），按 `AXXX` 标记 + manifest.txt
      还原成 369 个以原日文名命名的文件，输出到 split_out/。
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

TOOLS = Path(__file__).resolve().parent
SRC_DIR = TOOLS / "_texts_for_translation" / "characters"
# 合并产物单独放在 _merged_chars/，不进原流程目录，避免任何干扰
MERGED_DIR = TOOLS / "_merged_chars"
MERGED_FILE = MERGED_DIR / "characters_merged.txt"
MERGED_DOCX = MERGED_DIR / "characters_merged.docx"
MANIFEST_FILE = MERGED_DIR / "manifest.txt"

MARK_PREFIX = "A"  # 标记格式： A001  A002 ...（纯大写字母+数字，无单词，不会被翻译）


def fmt_idx(i: int) -> str:
    return f"{MARK_PREFIX}{i:03d}"


def is_marker(line: str) -> bool:
    s = line.strip()
    return s.startswith(MARK_PREFIX) and s[1:].isdigit()


def read_lines(path: Path) -> list[str]:
    """读取合并文件为行列表，支持 .txt 与 .docx。"""
    if path.suffix.lower() == ".docx":
        import docx
        doc = docx.Document(str(path))
        return [p.text for p in doc.paragraphs]
    return path.read_text(encoding="utf-8").splitlines()


def merge() -> int:
    files = sorted(SRC_DIR.glob("*.txt"))
    if not files:
        print(f"[ERROR] 未找到 {SRC_DIR}")
        return 1
    MERGED_DIR.mkdir(parents=True, exist_ok=True)

    # 本地映射：序号 -> 原日文文件名
    manifest_lines = [f"{fmt_idx(i)}<=>{f.name}" for i, f in enumerate(files, 1)]
    MANIFEST_FILE.write_text("\n".join(manifest_lines) + "\n", encoding="utf-8")

    parts: list[str] = []
    parts.append(
        "# characters 合并待翻译文件（共 %d 个）。\n"
        "# 翻译说明：\n"
        "#  - 每行开头的 [N] 序号请保留；只需把日文替换为中文译文。\n"
        "#  - 形如 A001 / A002 的标记行是文件分隔符，请勿改动、勿翻译。\n"
        "#  - 译完把本文件发回即可，会按 AXXX 标记拆回 369 个文件。\n" % len(files)
    )
    for i, f in enumerate(files, 1):
        parts.append(fmt_idx(i))
        parts.append(f.read_text(encoding="utf-8").rstrip("\n"))
        parts.append("")  # 文件间空行分隔
    MERGED_FILE.write_text("\n".join(parts).rstrip("\n") + "\n", encoding="utf-8")
    print(f"合并完成：{len(files)} 个文件 -> {MERGED_FILE}")
    print(f"映射表已存 -> {MANIFEST_FILE}")
    return 0


def to_docx() -> int:
    if not MERGED_FILE.is_file():
        print(f"[ERROR] 缺少 {MERGED_FILE}，先运行 merge")
        return 1
    import docx
    from docx.shared import RGBColor

    doc = docx.Document()
    for line in read_lines(MERGED_FILE):
        p = doc.add_paragraph(line)
        if is_marker(line):
            for run in p.runs:
                run.bold = True
        elif line.lstrip().startswith("#"):
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.save(str(MERGED_DOCX))
    print(f"docx 已生成 -> {MERGED_DOCX}（{len(read_lines(MERGED_FILE))} 段落）")
    return 0


def docx_to_txt(src: Path, dst: Path) -> int:
    lines = read_lines(Path(src))
    dst = Path(dst)
    dst.write_text("\n".join(lines).rstrip("\n") + "\n", encoding="utf-8")
    print(f"txt 已生成 -> {dst}（{len(lines)} 行）")
    return 0


def _src_indices(name: str) -> list[int]:
    p = SRC_DIR / name
    if not p.is_file():
        return []
    idx = []
    for line in p.read_text(encoding="utf-8").splitlines():
        m = re.match(r"^\[(\d+)\]", line)
        if m:
            idx.append(int(m.group(1)))
    return idx


def split(merged: Path, out_dir: Path | None = None) -> int:
    merged = Path(merged)
    if not merged.is_file():
        print(f"[ERROR] 文件不存在：{merged}")
        return 1
    if not MANIFEST_FILE.is_file():
        print(f"[ERROR] 缺少映射表 {MANIFEST_FILE}，无法还原原文件名")
        return 1
    name_map = {}
    for line in MANIFEST_FILE.read_text(encoding="utf-8").splitlines():
        if "<=>" not in line:
            continue
        key, name = line.split("<=>", 1)
        name_map[key.strip()] = name.strip()

    out_root = Path(out_dir) if out_dir else (MERGED_DIR / "split_out")
    out_root.mkdir(parents=True, exist_ok=True)

    # 先按 AXXX 切成 (key, 行列表)
    raw: dict[str, list[str]] = {}
    cur: str | None = None
    for line in read_lines(merged):
        s = line.strip()
        if is_marker(s):
            cur = s
            raw[cur] = []
        elif cur is not None:
            raw[cur].append(line)

    repaired = 0
    warned = 0
    for akey, lines in raw.items():
        fname = name_map.get(akey)
        if fname is None:
            warned += 1
            fname = f"__UNKNOWN_{akey}.txt"
        src_idx = _src_indices(fname)
        present = set()
        malformed = []
        for i, ln in enumerate(lines):
            m = re.match(r"^\[(\d+)\]", ln)
            if m:
                present.add(int(m.group(1)))
            elif ln.strip():
                malformed.append(i)
        missing = [n for n in src_idx if n not in present]
        # 修复：把错位/缺前缀的序号补回正确的 [N]
        for li in malformed:
            if not missing:
                warned += 1
                break
            mi = missing.pop(0)
            s = re.sub(r"^【\d+】\s*", "", lines[li])
            s = re.sub(r"^\[\d+\]\s*", "", s)
            lines[li] = f"[{mi}] {s}"
            repaired += 1
        if missing:
            warned += 1
            print(f"[WARN] {fname}: 仍有缺失序号 {missing}（译文未提供，将跳过注入）")
        (out_root / fname).write_text("\n".join(lines).rstrip("\n") + "\n", encoding="utf-8")

    print(f"拆分完成：{len(raw)} 个文件 -> {out_root}")
    if repaired:
        print(f"已自动修复错位/缺前缀的序号 {repaired} 处")
    if warned:
        print(f"[注意] {warned} 处需留意（见上方 WARN）")
    return 0


if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else "merge"
    if cmd == "split":
        target = Path(sys.argv[2]) if len(sys.argv) > 2 else MERGED_FILE
        out = Path(sys.argv[3]) if len(sys.argv) > 3 else None
        sys.exit(split(target, out))
    elif cmd == "docx":
        sys.exit(to_docx())
    elif cmd == "totxt":
        src = Path(sys.argv[2]) if len(sys.argv) > 2 else MERGED_DIR / "characters_merged_translated.docx"
        dst = Path(sys.argv[3]) if len(sys.argv) > 3 else MERGED_DIR / "characters_merged_translated.txt"
        sys.exit(docx_to_txt(src, dst))
    else:
        sys.exit(merge())
